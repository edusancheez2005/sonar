from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


ACCENT = RGBColor(0, 229, 255)
PRIMARY = RGBColor(54, 166, 186)
BODY_COLOR = RGBColor(50, 50, 50)
MUTED = RGBColor(120, 120, 120)


def set_run(run, size=10, bold=False, italic=False, color=BODY_COLOR, font_name="Calibri"):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.name = font_name


def add_title(doc, text, size=22, color=ACCENT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=size, bold=True, color=color)
    p.space_after = Pt(4)
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f">  {text}")
    set_run(r, size=14, bold=True, color=ACCENT)
    p.space_after = Pt(6)
    return p


def add_sub_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=11, bold=True, color=PRIMARY)
    p.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=10, color=BODY_COLOR)
    p.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r, size=10, color=BODY_COLOR)
    p.space_after = Pt(2)
    return p


def add_qa(doc, question, answer):
    p = doc.add_paragraph()
    r = p.add_run(f"Q: {question}")
    set_run(r, size=10, bold=True, color=PRIMARY)
    p.space_after = Pt(2)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(answer)
    set_run(r2, size=10, color=BODY_COLOR)
    p2.space_after = Pt(10)
    return p2


def add_tip(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"TIP: {text}")
    set_run(r, size=9.5, italic=True, color=PRIMARY)
    p.space_after = Pt(8)
    return p


def generate():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_path = os.path.join(base, "ACCELERATOR_INTERVIEW_PREP.docx")

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    # ============================
    # COVER
    # ============================
    add_title(doc, "King's Start-up Accelerator", size=24, color=ACCENT)
    add_title(doc, "Interview Prep Guide", size=18, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Eduardo Sanchez Morales  |  Sonar\nApril 2026")
    set_run(r, size=11, color=MUTED)
    p.space_after = Pt(12)

    add_body(doc,
        "This guide covers your 2-3 minute pitch script, prepared answers for "
        "Andrea's specific feedback questions, anticipated panel questions, and "
        "key metrics and talking points to keep front of mind."
    )

    doc.add_page_break()

    # ============================
    # INTERVIEW FORMAT
    # ============================
    add_section_title(doc, "Interview Format & Logistics")
    add_bullet(doc, "Virtual on MS Teams, 30 minutes sharp")
    add_bullet(doc, "Panel of 3: Andrea + ventures/incubations team member + external network member")
    add_bullet(doc, "2-3 min verbal pitch (NO slides) followed by Q&A")
    add_bullet(doc, "You can ask them questions at the end")
    add_bullet(doc, "They want to see: innovation, validation, and founder-market fit")
    add_bullet(doc, "Timeline: interviews 13-23 April, offers early May, welcome week 17-19 June")

    add_tip(doc,
        "The pitch brief asks for 4 things: Unmet Need, Your Solution, Competitive Advantage, "
        "and Team/Motivation. Structure your 2-3 min around these exact pillars."
    )

    # ============================
    # PITCH SCRIPT
    # ============================
    add_section_title(doc, "Your 2-3 Minute Pitch Script")
    add_body(doc,
        "Below is a structured script hitting all four areas from the brief. "
        "Practice it out loud. Aim for 2 minutes 15 seconds to leave buffer. "
        "Speak naturally, don't memorise word for word, but nail the structure and key numbers."
    )

    add_sub_title(doc, "1. Unmet Need & Customer (30-40 sec)")
    add_body(doc,
        "Crypto traders lose money every day, not because the market is random, but because "
        "the information they need is either too expensive, scattered across too many tools, or "
        "just unreliable. I lived this myself as a trader. I was constantly switching between Whale Alert "
        "for transaction data, LunarCrush for sentiment, CoinGecko for prices, and Twitter for news. "
        "I never had the full picture. When a whale moved $10 million of Ethereum to an exchange, "
        "I'd find out hours later. By then the price had already moved.\n\n"
        "Through Idea Factory I spoke with over 20 active traders and the pattern was clear. They're "
        "semi-experienced, 1 to 3 years in crypto, hold portfolios between 1K and 100K, and they're "
        "frustrated by fragmented tools. Nansen costs 150 a month. Dune requires SQL. Whale Alert "
        "gives raw data with zero context. No single platform brings it all together affordably."
    )

    add_sub_title(doc, "2. The Solution (30-40 sec)")
    add_body(doc,
        "That's why I built Sonar. It's an AI-powered crypto intelligence platform that combines "
        "whale tracking, multi-source sentiment analysis, news, and a conversational AI advisor "
        "called Orca, all in one place for 7.99 a month.\n\n"
        "We track real-time whale transactions above $500K across major blockchains, run a "
        "10-component sentiment algorithm fusing 6 data sources into a single confidence-scored "
        "rating, and let traders ask Orca natural language questions about any token. We're also "
        "developing a backtesting engine so traders can validate our signals against historical "
        "data before risking real money. And longer term, we're building towards a proprietary "
        "LLM trained on our own signal data and market history, essentially an AI quant analyst "
        "that gets smarter the more data we feed it."
    )

    add_sub_title(doc, "3. Competitive Advantage (20-30 sec)")
    add_body(doc,
        "No other platform combines whale tracking, sentiment, news, social intelligence, and "
        "an AI advisor at this price point. Our competitors solve pieces of the puzzle. Nansen does "
        "on-chain but charges 150 a month and has no AI. Arkham does blockchain intelligence but "
        "it's built for researchers, not everyday traders. LunarCrush does social but ignores whale "
        "activity entirely. Sonar is the only tool that gives traders the complete picture for under "
        "8 pounds a month. And our long-term moat is the proprietary LLM we're building, trained "
        "on our own accumulated signal data. Think of it like a hedge fund quant, but accessible "
        "to retail traders."
    )

    add_sub_title(doc, "4. Team & Traction (20-30 sec)")
    add_body(doc,
        "I'm a CS with Management student at King's with extensive AI experience, including an "
        "AI internship in the UAE. I built the entire platform from scratch. My co-founder Saif "
        "is a CS and Maths student, also at King's, and he leads our Bitcoin and Solana data "
        "engineering. Between us we cover full-stack, AI, and blockchain development without "
        "needing to hire.\n\n"
        "The product is live at sonartracker.io with over 650 users, and we grew 150 of those "
        "in just the last 2 weeks with zero marketing spend. We have 5 paying subscribers and "
        "Stripe is processing real payments. We deliberately made the free version more accessible "
        "to maximise the top of funnel, then we'll convert users to premium as they see the value. "
        "6 automated data pipelines run in production across 100+ tokens.\n\n"
        "The impact I want to make is simple: democratise access to institutional-grade crypto "
        "analytics. Right now the best tools are only available to people who can afford hundreds "
        "a month. Sonar makes that same data accessible to everyone."
    )

    add_tip(doc,
        "Aim for confident, conversational delivery. The panel wants a snapshot, not a sales pitch. "
        "End with traction numbers, they're your strongest proof point."
    )

    doc.add_page_break()

    # ============================
    # ANDREA'S FEEDBACK
    # ============================
    add_section_title(doc, "Answering Andrea's Specific Feedback")
    add_body(doc,
        "These are the exact areas Andrea flagged in her email. The panel WILL ask about these. "
        "Have these answers locked in."
    )

    add_qa(doc,
        "Structured validation - what hypothesis are you testing and what metrics are you tracking?",
        "We're testing four core hypotheses right now. First, that combining multiple data sources "
        "into one composite score genuinely improves trading outcomes versus using individual tools. "
        "We're developing our backtesting engine specifically to test this, tracking hit rates and "
        "returns against buy-and-hold baselines. Second, that retail traders will pay 7.99 a month "
        "when free but fragmented alternatives exist. We deliberately made the free tier more "
        "accessible to grow the user base first, and we're measuring conversion rate as we start "
        "nudging users towards premium. We have 5 paying subscribers from 650+ users. Third, that "
        "an AI chat interface is the most valued feature. We track Orca AI conversations per user "
        "per week. Fourth, that users retain week over week. We track weekly active user retention. "
        "Since the application, I've become much more structured about logging these metrics and "
        "making decisions based on them rather than gut feel."
    )

    add_qa(doc,
        "The pricing gap: have you estimated the value of your product to users? Is the assumption "
        "that retail traders are priced out dependent on portfolio size?",
        "This is a really good challenge and I've thought about it more since the application. "
        "You're right that a trader with a 100K portfolio might happily pay 150 a month for Nansen "
        "if it helps them make even 1% better returns. The cost-to-portfolio ratio matters.\n\n"
        "But our research shows the sweet spot is traders with 1K to 50K portfolios, which is the "
        "vast majority of retail. For someone with a 5K portfolio, 150 a month is 3% of their entire "
        "holdings just on analytics, which makes no sense. At 7.99 that's 0.16%, which is very "
        "reasonable.\n\n"
        "That said, there is room for tiered pricing as we grow. A Pro tier at 19.99 or 29.99 for "
        "larger portfolio traders who want deeper analytics, API access, or higher AI usage limits. "
        "We intentionally started at 7.99 to maximise adoption and learn, but the pricing model "
        "will evolve as we understand willingness to pay at different portfolio sizes."
    )

    add_qa(doc,
        "Could Nansen develop a retail version and be a threat? How do you build a sustainable "
        "competitive advantage?",
        "Nansen absolutely could launch a cheaper tier, and I'd expect them to eventually. But there "
        "are a few reasons I think Sonar still wins in that scenario.\n\n"
        "First, Nansen's entire business model is built around high-value enterprise clients. Their "
        "cost structure, sales team, and product design are optimised for institutions paying thousands "
        "a year. Shipping a 7.99 retail product would cannibalise their premium pricing and confuse "
        "their positioning. It's the classic innovator's dilemma. Same applies to Arkham, they do "
        "blockchain intelligence with entity tracking but their interface is built for researchers "
        "and analysts, not everyday retail traders. No sentiment, no AI chat.\n\n"
        "Second, Nansen and Arkham both focus on on-chain analytics. They don't do sentiment analysis, "
        "they don't do news NLP, they don't have a conversational AI advisor. Sonar's advantage is "
        "the combination, not any single feature.\n\n"
        "Third, our moat deepens over time. We're building towards a proprietary LLM trained on our "
        "own accumulated signal data, whale movement history, and market outcomes. Think of it like "
        "building an AI quant that gets smarter every day. No competitor is doing this at the retail "
        "level. We also plan to build our own blockchain data pipeline, reducing dependency on "
        "third-party APIs and giving us data depth that's hard to replicate.\n\n"
        "The sustainable competitive advantage is: first mover in the affordable all-in-one space, "
        "a proprietary AI that compounds in intelligence over time, and a community of engaged "
        "retail traders that we're building now."
    )

    add_qa(doc,
        "Your commitment to the in-person programme and community?",
        "I'm fully committed. I've already experienced the value of the EI community through Idea "
        "Factory and Open Pitch Night, both were genuinely useful for shaping Sonar. I plan to be "
        "in the co-working space regularly, I'll actively participate in workshops, and I want to "
        "contribute back. On the technical side I can help other founders with web dev, AI integration, "
        "API design, and shipping to production. I'm happy to run informal sessions or pair-code "
        "with anyone who needs help. I'll show up."
    )

    doc.add_page_break()

    # ============================
    # ANTICIPATED QUESTIONS
    # ============================
    add_section_title(doc, "Anticipated Panel Questions")
    add_body(doc,
        "Beyond Andrea's feedback areas, here are questions the panel is likely to ask based on "
        "the brief's focus on innovation, validation, and founder-market fit."
    )

    add_qa(doc,
        "What's your riskiest assumption right now?",
        "That combining multiple data sources into one composite score genuinely improves trading "
        "outcomes versus using individual tools. If our signals aren't measurably better, the value "
        "proposition weakens. That's exactly why we're developing the backtesting engine, so traders "
        "can test this with real data and we can iterate on the algorithm. We want hard evidence "
        "before making strong claims, not just vibes."
    )

    add_qa(doc,
        "Your conversion rate is low. Is that good enough? What's your plan to improve it?",
        "So right now the conversion rate looks low but that's intentional. We deliberately made "
        "the free version more accessible recently to maximise the top of funnel, we grew 150 users "
        "in just 2 weeks with no marketing at all. The strategy is to get as many traders in as "
        "possible, let them experience the value, and then convert them gradually. The industry "
        "benchmark for freemium SaaS is around 2-5% and we'll get there as we start activating "
        "conversion. Our plan is threefold. First, improve the onboarding flow so new users "
        "immediately see value. Second, use Orca AI as the conversion hook because it's the feature "
        "traders get most excited about and we limit it to 1 question a day on free. Third, "
        "implement email nurture sequences that educate free users about premium features they're "
        "missing."
    )

    add_qa(doc,
        "How do you acquire users? What's your growth strategy?",
        "Right now growth is organic through crypto communities on Twitter/X, Reddit, and Discord. "
        "We're also working on a content marketing strategy: short-form video series explaining crypto "
        "concepts in an animated format, published on TikTok, Instagram Reels, and YouTube Shorts. "
        "A friend doing a similar format for history content hit 30K followers in 2 weeks, so the "
        "format is validated. We've designed bridge episodes that specifically funnel viewers to Sonar. "
        "Beyond that, SEO through our blog and glossary pages, and community building through a "
        "Sonar Discord server."
    )

    add_qa(doc,
        "What would you do with 10x more users tomorrow? Can your infrastructure handle it?",
        "Yes. The entire platform runs on Vercel's serverless infrastructure and Supabase, both of "
        "which scale automatically. Our data pipelines are independent cron jobs, they don't depend "
        "on user load. The only bottleneck would be AI API costs if Orca usage spiked, but we manage "
        "that through daily usage limits on both free and premium tiers. 10x users would actually be "
        "great because our unit economics improve with scale, fixed pipeline costs are spread across "
        "more subscribers."
    )

    add_qa(doc,
        "What have you learned since submitting the application?",
        "A lot actually. First, onboarding matters more than I thought. Users who don't understand "
        "the dashboard in the first 30 seconds tend to bounce, so we've been simplifying the initial "
        "experience. Second, Orca AI conversations are stickier than I expected. Users who try Orca "
        "come back more often, which confirms it should be the centrepiece of the premium offering. "
        "Third, growth can happen fast even without marketing. We went from 400 to 650+ users "
        "organically, 150 of those in just the last 2 weeks. That tells me the product has real "
        "pull. Fourth, I've gotten more structured about tracking metrics weekly rather than just "
        "building features and hoping for the best."
    )

    add_qa(doc,
        "Why should we pick you over other applicants?",
        "Because Sonar isn't an idea, it's a live product with 650+ users and real revenue. I've "
        "built the entire platform from scratch. I'm a CS student with deep AI experience, I did "
        "an AI internship in the UAE before starting this, so I have the technical foundation to "
        "build and ship fast without burning money on developers. My co-founder Saif is a CS and "
        "Maths student at King's handling our blockchain data engineering. Between us we move fast. "
        "What I need from the accelerator is the business mentorship, network, and structured "
        "accountability to match the speed of execution I already have on the technical side."
    )

    add_qa(doc,
        "What's your biggest challenge right now?",
        "User acquisition. The product works, the data is solid, and early users are engaged. But "
        "getting in front of enough crypto traders consistently is the bottleneck. I'm a strong "
        "builder but I need to become a stronger marketer. That's one of the things I'm hoping the "
        "accelerator can help with, both through mentorship and the network."
    )

    add_qa(doc,
        "How do you handle the regulatory landscape for crypto in the UK?",
        "We don't provide financial advice and we're very clear about that. Sonar provides data and "
        "analysis, not recommendations to buy or sell. Every signal includes a confidence score so "
        "users understand the uncertainty. We're also mindful of ASA and FCA guidelines around crypto "
        "marketing, so our content focuses on education and data transparency rather than promising "
        "returns. As the regulatory landscape evolves we'll adapt, but our core product is analytics, "
        "not trading, which keeps us in a safer position."
    )

    add_qa(doc,
        "What does success look like in 12 months?",
        "5,000 active users, 250+ paying subscribers, meaningful monthly recurring revenue that "
        "covers costs with margin, a launched public API, and a clear path to pre-seed investment. "
        "But the metric I care about most is whether traders using Sonar are actually making better "
        "decisions. If our backtesting data shows our signals consistently outperform random, that's "
        "the real proof that we've built something valuable."
    )

    add_qa(doc,
        "How do you balance studies and the startup?",
        "I'm disciplined about time. I run weekly sprints with clear priorities, and my co-founder "
        "Saif shares the workload, he's a CS and Maths student at King's and leads our Bitcoin and "
        "Solana data engineering. I'm in my final two years of CS with Management, and honestly "
        "the degree and the startup feed into each other. The management modules help me think "
        "about positioning and growth, and the CS modules keep my technical skills sharp. Having "
        "done an AI internship in the UAE before this, I also know how to structure my time around "
        "intense technical work. I make sure to exercise, take time offline, and lean on my support "
        "network when things get intense."
    )

    doc.add_page_break()

    # ============================
    # KEY METRICS CHEAT SHEET
    # ============================
    add_section_title(doc, "Key Numbers & Talking Points")
    add_body(doc, "Keep these front of mind. Drop them naturally in conversation.")

    add_sub_title(doc, "Traction")
    add_bullet(doc, "650+ users (grew 150 in last 2 weeks with zero marketing)")
    add_bullet(doc, "5 paying subscribers at 7.99/month")
    add_bullet(doc, "480 ARR (annual recurring revenue)")
    add_bullet(doc, "Product live at sonartracker.io since early 2026")
    add_bullet(doc, "Stripe processing real payments")
    add_bullet(doc, "100+ tokens tracked across major blockchains")
    add_bullet(doc, "6 automated data pipelines running in production")
    add_bullet(doc, "Free tier deliberately opened up to maximise top of funnel")

    add_sub_title(doc, "Unit Economics")
    add_bullet(doc, "Operating costs: ~500/month (APIs, Vercel, Supabase, AI inference)")
    add_bullet(doc, "Premium: 7.99/month = 95.88/year per subscriber")
    add_bullet(doc, "Break-even at roughly 63 subscribers at current cost base")
    add_bullet(doc, "Future tiers: Pro at 19.99-29.99, API tier, institutional tier")
    add_bullet(doc, "Costs will decrease as we build proprietary data pipelines and our own LLM")

    add_sub_title(doc, "Market")
    add_bullet(doc, "420M+ crypto holders globally, 50-80M active traders")
    add_bullet(doc, "Crypto analytics market: $2.2B (2024) projected to $8.6B by 2030 (25% CAGR)")
    add_bullet(doc, "20-30M English-speaking active traders (initial focus)")
    add_bullet(doc, "Nansen raised $75M, Dune reached $1B valuation, Arkham launched own token")

    add_sub_title(doc, "Competitive Positioning")
    add_bullet(doc, "Nansen: $150/month, no AI, no sentiment, enterprise focus")
    add_bullet(doc, "Arkham: blockchain intelligence but complex UI, no sentiment, no AI chat, researcher-focused")
    add_bullet(doc, "Glassnode: $40-800/month, no AI, limited altcoins")
    add_bullet(doc, "Whale Alert: free but raw data, no analysis or context")
    add_bullet(doc, "LunarCrush: social only, ignores on-chain whale activity")
    add_bullet(doc, "Dune: powerful but requires SQL, inaccessible to most retail")
    add_bullet(doc, "Sonar: ONLY platform combining all + AI advisor for 7.99/month")
    add_bullet(doc, "Long-term moat: proprietary LLM trained on our own signal data = AI quant for retail")

    add_sub_title(doc, "Grants & Validation")
    add_bullet(doc, "2,000 from Idea Factory + 200 from Open Pitch Night")
    add_bullet(doc, "Customer discovery: 20+ trader interviews through Idea Factory")
    add_bullet(doc, "Applying for Solana Foundation, Ethereum Foundation, Polygon grants")

    # ============================
    # QUESTIONS TO ASK
    # ============================
    add_section_title(doc, "Questions to Ask the Panel")
    add_body(doc,
        "Have 2-3 ready. Asking good questions shows you're serious about the programme, "
        "not just the badge."
    )
    add_bullet(doc,
        "What does the most successful founder from a previous cohort have in common in terms "
        "of how they used the programme?"
    )
    add_bullet(doc,
        "Are there mentors in the network with experience in fintech, crypto, or data products "
        "specifically?"
    )
    add_bullet(doc,
        "How structured is the workshop programme versus self-directed work time? I want to "
        "make sure I'm balancing building with learning."
    )
    add_bullet(doc,
        "What's the biggest mistake you see accelerator founders make in the first 3 months?"
    )

    # ============================
    # FINAL TIPS
    # ============================
    add_section_title(doc, "Final Tips")
    add_bullet(doc, "Be conversational, not rehearsed. They want to meet Eduardo, not a pitch bot.")
    add_bullet(doc, "Lead with the problem and your personal experience. It's your strongest hook.")
    add_bullet(doc, "When you don't know something, say so honestly and explain how you'd find out.")
    add_bullet(doc, "Numbers build credibility. Drop traction stats naturally throughout your answers.")
    add_bullet(doc, "Show coachability. Mention what you've learned, what you got wrong, what you changed.")
    add_bullet(doc, "End strong. Your closing impression matters. Be enthusiastic about the programme specifically.")
    add_bullet(doc, "Test your Teams setup beforehand. Camera on, good lighting, clean background.")
    add_bullet(doc, "Have sonartracker.io open in a tab in case they ask for a quick demo walkthrough.")

    doc.save(output_path)
    print(f"DOCX generated at: {output_path}")


if __name__ == "__main__":
    generate()
