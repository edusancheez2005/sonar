"""Generate professional Word document - Sonar Meeting Prep v2."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO = os.path.join(BASE, "public", "assets", "logo.png")
OUTPUT = os.path.join(BASE, "Sonar_Meeting_Prep_v2.docx")
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2.5); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11); style.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
for lvl, sz in [('Heading 1',22),('Heading 2',16),('Heading 3',13)]:
    s = doc.styles[lvl]; s.font.size = Pt(sz); s.font.color.rgb = RGBColor(0x0a,0x6e,0x82); s.font.bold = True

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; p.add_run(f" {text}")
    else: p.add_run(text)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

# Footer
for s in doc.sections:
    f = s.footer; f.is_linked_to_previous = False
    fp = f.paragraphs[0] if f.paragraphs else f.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO): fp.add_run().add_picture(LOGO, width=Inches(0.6))
    r = fp.add_run("  Sonar Tracker  |  Confidential  |  March 2026"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x88,0x88,0x88)

# ═══ TITLE PAGE ═══
if os.path.exists(LOGO):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(LOGO, width=Inches(1.5))
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = t.add_run("Sonar Tracker"); r.font.size = Pt(32); r.font.color.rgb = RGBColor(0x0a,0x6e,0x82); r.bold = True
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = s.add_run("Cofounder Meeting Prep\nMarch 2026"); r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x5a,0x6a,0x7a)
doc.add_paragraph(); doc.add_paragraph()
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = d.add_run("CONFIDENTIAL"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xaa,0xaa,0xaa); r.italic = True
doc.add_page_break()

# ═══ 1. TEAM ALIGNMENT ═══
doc.add_heading("Team Alignment", level=1)
p = doc.add_paragraph(); p.add_run("This is not a side project anymore. ").bold = True
p.add_run("We have 500 users, grant applications in progress, and a pipeline processing real whale transactions across 7 chains. This is a company.")
doc.add_paragraph()
for b in ["Minimum 10-15 hours/week each. Non-negotiable.", "Every sprint is 1 week. Clear deliverables. Friday deadline.", "Weekly call: review what shipped, plan next week.", "Ship small things frequently, not big things rarely.", "If something is blocking you, say it immediately. Don't wait."]:
    bullet(b)

doc.add_page_break()

# ═══ 2. COMPETITOR ANALYSIS ═══
doc.add_heading("Competitor Gap Analysis", level=1)

doc.add_heading("Arkham Intelligence ($12M+ raised)", level=2)
table(["Feature","Arkham","Sonar","Gap"],[
    ["Entity labels (100K+ wallets labelled)","Yes","Partial","Need label DB"],
    ["Real-time transfer feed","Live stream","Live ticker","Even"],
    ["Exchange flow tracking (net in/out)","Dedicated view","Dashboard","Even"],
    ["Wallet profiling + PnL tracking","Deep profiles","No","MAJOR GAP"],
    ["AI-generated trending insights","Auto bulletins","No","Need proactive AI"],
    ["Custom alerts (wallet/token)","Yes","No","MAJOR GAP (Week 3-4)"],
    ["Token holder analysis","Yes","Partial","Medium gap"],
    ["Developer API","Full API","No","Week 10"],
    ["Twitter/X influencer tracking","No","No","OPPORTUNITY"],
    ["Conversational AI advisor","No","ORCA (live)","OUR ADVANTAGE"],
])

doc.add_heading("Nansen ($75M+ raised)", level=2)
table(["Feature","Nansen","Sonar","Gap"],[
    ["Smart Money tracking (labelled)","Core product","70K wallets (basic)","Need better labels"],
    ["Smart Alerts (push notifications)","Killer feature","No","MAJOR GAP (Week 3-4)"],
    ["Agentic Trading (AI executes)","Yes (new, mobile)","No","Future (Week 8)"],
    ["18+ chain support","Yes","7 chains","Expanding"],
    ["Token God Mode (deep analytics)","Yes","Partial token page","Improving"],
    ["CLI / MCP for AI agents","Yes (new)","No","Future"],
])

doc.add_heading("What Sonar Has That NOBODY Else Has", level=2)
table(["Feature","Why It Matters"],[
    ["ORCA AI with live multi-source data injection","Nobody answers trader questions with real-time whale + sentiment + price data"],
    ["4-tier composite signal engine","Unique: whale flow + momentum + sentiment + community fused into one score"],
    ["Trap detection (bullish traps, dead-cat bounces)","Catches dangerous patterns competitors miss"],
    ["Completely free platform","$0 vs $50-150/mo. Massive accessibility advantage"],
    ["AI buy/sell classification (not just raw transfers)","We tell you the intent, not just the transfer"],
])

doc.add_heading("Opportunities Nobody Has Built Yet", level=2)
doc.add_paragraph()
bullet("Twitter/X Influencer Whale Correlation: Track Elon Musk, Donald Trump, CZ, Vitalik, top CT accounts. When they tweet about a token, correlate with whale movement data in real time. 'Elon tweeted about DOGE 15 min ago. Whale bought $2M DOGE 5 min later.' Nobody connects influence to on-chain flow.", bold_prefix="1.")
bullet("AI-Powered Copy-Trading Bot: Mirror whale wallet movements automatically. Users pick a smart money wallet, bot paper-trades (then real-trades) based on their moves. Arkham has intel but no execution. We combine both.", bold_prefix="2.")
bullet("Predictive Whale Alerts: ORCA doesn't wait for you to ask. 'A whale moved $10M SOL to Binance. Historically, this precedes a 5-8% drop 70% of the time.' Predictive, not reactive. Nobody does this.", bold_prefix="3.")
bullet("Token Risk Score (0-100): One number fusing whale concentration + social sentiment + price volatility + holder distribution. 'Is this safe to buy right now?' in one glance. Nobody does this well.", bold_prefix="4.")

doc.add_page_break()

# ═══ 3. VALUE CHAIN ═══
doc.add_heading("The Value Chain", level=1)
table(["Stage","What","Status"],[
    ["1. Data Collection","Raw blockchain transactions across chains","7 chains live. Done."],
    ["2. Data Classification","Classify BUY/SELL/TRANSFER with confidence","ERC-20 strong, SOL/BTC basic"],
    ["3. Data Analysis","Signal engine, sentiment, net flows","Live for tracked tokens"],
    ["4. AI Intelligence","ORCA answers questions with live data","Live, reactive only"],
    ["5. Social Intelligence","Track influencer tweets, correlate with chain","NOT STARTED"],
    ["6. Alerts + Notifications","Push when whales move, patterns form","NOT STARTED"],
    ["7. Custom AI Model","Fine-tune on classified Sonar data","NOT STARTED"],
    ["8. Autonomous Agent","24/7 monitoring, proactive insights","CONCEPT"],
])

doc.add_page_break()

# ═══ 4. SAIF'S ROADMAP (12 WEEKS) ═══
doc.add_heading("Saif: 12-Week Dev Roadmap", level=1)

weeks = [
    ("Week 1: Dashboard + Data Fix", [
        "Verify Solana + Bitcoin transactions display correctly on dashboard with proper labels and classification",
        "Fix dashboard loading speed: add DB indexes, cache /api/dashboard/summary, target sub-2-second load",
        "Expand Solana token coverage: add RNDR, MOBILE, JITO, MSOL, BSOL, HNT, WEN, and top 20 SPL tokens",
        "Add CryptoPanic as secondary news source alongside existing pipeline",
    ]),
    ("Week 2: Solana Classification", [
        "Improve Solana classification: replace rough USD estimate with actual CoinGecko price at transaction time",
        "Map known Solana CEX deposit addresses (Binance, Coinbase, Kraken, OKX, Bybit)",
        "Add Jupiter swap signature detection and Raydium pool interaction parsing for better buy/sell accuracy",
    ]),
    ("Week 3: Alerts Backend", [
        "Design alert system: Supabase user_alerts table (trigger_type, token, threshold, direction, delivery_method)",
        "Build alert trigger engine: evaluate new transactions against user alert rules every 60 seconds",
        "Alert types: whale moved $X of TOKEN, net flow reversal, unusual volume spike, watched wallet activity",
    ]),
    ("Week 4: Alerts Delivery + UI", [
        "Email notification delivery via Resend for triggered alerts",
        "In-app notification bell on navbar showing recent triggered alerts",
        "Alert creation UI: users set up custom alerts (pick token, set threshold, choose direction)",
        "Test end-to-end: whale moves SOL > user gets email + in-app notification within 2 minutes",
    ]),
    ("Week 5: Wallet Intelligence", [
        "Build exchange address database: map known CEX addresses across ETH, SOL, BTC, Polygon, Tron",
        "Entity labelling: map top 500 whale addresses to known entities (funds, exchanges, OG whales, protocols)",
        "Wallet profile page: transaction history, net position, PnL estimate, classification breakdown",
    ]),
    ("Week 6: Twitter/X Influencer Tracking", [
        "Build Twitter/X monitoring pipeline: poll or stream tweets from 50+ curated crypto influencer accounts",
        "Key accounts: Elon Musk, Donald Trump, CZ, Vitalik, Brian Armstrong, Michael Saylor, top CT accounts",
        "Extract token mentions from tweets (ticker detection + NLP)",
        "Correlate: when influencer tweets about TOKEN, check whale movements within +-30 minutes",
        "Surface on dashboard: 'Elon tweeted about DOGE. 15 min later, whale bought $2M DOGE on-chain.'",
    ]),
    ("Week 7: Proactive AI Insights", [
        "Auto-generated daily whale report: top 5 accumulations, top 5 distributions, biggest single moves",
        "ORCA proactive mode: hourly scan, generate insight when unusual pattern detected (volume spike, flow reversal)",
        "Historical pattern matching: 'This whale-to-exchange pattern preceded a 5-8% drop in 3 of last 5 similar cases'",
    ]),
    ("Week 8: Trading Bot Prototype", [
        "Copy-trading architecture: user selects a whale wallet to mirror",
        "Paper trading mode: simulate trades based on whale signals, track hypothetical PnL over 30 days",
        "Backtest engine: run historical whale signals against actual price outcomes to measure signal accuracy",
    ]),
    ("Week 9: Token Risk Score", [
        "Design composite risk score (0-100): whale concentration, social sentiment, volatility, holder distribution",
        "Scoring engine: calculate risk score for every tracked token in real time",
        "UI: risk badge on token pages and dashboard (green safe / amber caution / red high risk)",
    ]),
    ("Week 10: Public API v1", [
        "API endpoints: /api/v1/whale-transactions, /api/v1/signals, /api/v1/sentiment, /api/v1/alerts",
        "API key management and rate limiting",
        "Documentation page with curl examples and integration guide",
    ]),
    ("Week 11: Fine-Tuning Data Pipeline", [
        "Training dataset format: classified transaction + price outcome at T+1h, T+4h, T+24h",
        "Export pipeline: generate training CSVs from whale_transactions table with price joins",
        "Research Azure ML / Foundry / Hugging Face for fine-tuning costs and infrastructure",
        "Target: 200K+ labelled training examples ready",
    ]),
    ("Week 12: Model Fine-Tuning Prototype", [
        "Fine-tune GPT-4o-mini or Llama 3 8B on Sonar classified data",
        "Evaluate: does fine-tuned model classify more accurately than rule engine?",
        "If yes: integrate as ORCA's primary classification backend",
        "If no: iterate on training data quality, add more features",
    ]),
]

for title, items in weeks:
    doc.add_heading(title, level=2)
    for item in items:
        bullet(item)

doc.add_page_break()

# ═══ 5. EDUARDO'S ROADMAP ═══
doc.add_heading("Eduardo: 12-Week Founder Roadmap", level=1)

edu_weeks = [
    ("Week 1: Grant Submissions", [
        "Submit Solana Foundation grant ($15K)",
        "Submit Superteam UK microgrant ($10K)",
        "Fix Google Search Console image error (done)",
    ]),
    ("Week 2: Marketing Foundation", [
        "Set up content automation (Anthropic + OpenClaw for blog/social)",
        "First marketing video (Holo AI)",
        "OpenAI image generation for social media assets",
    ]),
    ("Week 3: Ask ORCA Everywhere", [
        "'Ask ORCA' contextual buttons on every dashboard section, token page, whale feed",
        "One-click: see a whale move, ask ORCA what it means",
    ]),
    ("Week 4: Daily Whale Report", [
        "Auto-generated daily whale activity summary (top moves, accumulations, distributions)",
        "Email delivery to all registered users",
        "Shareable public link per report (massive SEO value)",
    ]),
    ("Week 5: Event Prep", [
        "6-slide pitch deck for investor meetings",
        "Prepare for IFGS (April 21, London) and TOKEN2049 Dubai (April 29-30)",
        "Book 10+ investor meetings around events",
    ]),
    ("Week 6: More Grants + Accelerators", [
        "Apply to Barclays Eagle Labs, Techstars London, Entrepreneurs First (deadline April 6)",
        "Submit Base Builder Grants, Arbitrum Foundation, BNB Chain grants",
    ]),
    ("Week 7: SEO + Content Push", [
        "5 targeted long-form blog posts per week",
        "Solana whale tracking landing page",
        "Competitor comparison pages: 'Sonar vs Nansen', 'Sonar vs Arkham'",
    ]),
    ("Week 8: Community Building", [
        "Launch Sonar Discord / Telegram community",
        "Weekly whale recap thread on Twitter/X",
        "Engage Solana community channels with real whale insights",
    ]),
    ("Week 9: Landing Page + Pricing Update", [
        "Update all copy to reflect free model",
        "Feature Twitter influencer tracking as new capability",
        "Add social proof / testimonials section",
    ]),
    ("Week 10: Investor Materials", [
        "One-pager for angel investors (SEIS/EIS ready)",
        "Financial model: unit economics, growth projections, runway",
        "Build CRM: 30 UK angels + 15 crypto funds pipeline",
    ]),
    ("Week 11: AI Strategy", [
        "Research Foundry / Azure ML / Hugging Face fine-tuning costs",
        "Design training pipeline with Saif",
        "Define evaluation metrics: classification accuracy, signal accuracy vs price outcome",
    ]),
    ("Week 12: Partnerships", [
        "Reach out to Helius for partnership / sponsorship",
        "Reach out to LunarCrush for data partnership upgrade",
        "Evaluate: deeper analysis on existing chains vs more chain coverage",
    ]),
]

for title, items in edu_weeks:
    doc.add_heading(title, level=2)
    for item in items:
        bullet(item)

doc.add_page_break()

# ═══ 6. CORE FEATURES ═══
doc.add_heading("Core Feature Priority", level=1)

doc.add_heading("MUST BUILD (Weeks 1-4)", level=2)
p = doc.add_paragraph(); p.add_run("The basics our competitors already have:").italic = True
bullet("Real-time whale alerts with push notifications (email + in-app)", bold_prefix="1.")
bullet("Dashboard sub-2-second load time", bold_prefix="2.")
bullet("Solana/BTC transactions visible with proper classification", bold_prefix="3.")
bullet("20+ Solana SPL tokens tracked", bold_prefix="4.")

doc.add_heading("SHOULD BUILD (Weeks 5-8)", level=2)
p = doc.add_paragraph(); p.add_run("What makes people recommend Sonar to friends:").italic = True
bullet("Twitter/X influencer tracking correlated with whale movements", bold_prefix="5.")
bullet("Ask ORCA contextual buttons everywhere in the UI", bold_prefix="6.")
bullet("Wallet profiles with entity labels and PnL", bold_prefix="7.")
bullet("Daily auto-generated whale report (email + shareable)", bold_prefix="8.")
bullet("Trading bot prototype (paper trading based on whale signals)", bold_prefix="9.")

doc.add_heading("WANT TO BUILD (Weeks 9-12)", level=2)
p = doc.add_paragraph(); p.add_run("What makes Sonar unique in the industry:").italic = True
bullet("Token Risk Score (0-100 composite)", bold_prefix="10.")
bullet("Public API for developers", bold_prefix="11.")
bullet("Fine-tuned Sonar AI model on proprietary data", bold_prefix="12.")
bullet("Autonomous ORCA agent (24/7 monitoring + proactive alerts)", bold_prefix="13.")

doc.add_page_break()

# ═══ 7. METRICS ═══
doc.add_heading("Key Metric Targets", level=1)
table(["Metric","Now","4 Weeks","8 Weeks","12 Weeks"],[
    ["Users","500","800","1,500","3,000"],
    ["DAU","~20","50","150","300"],
    ["ORCA conversations/day","~5","15","50","100"],
    ["Chains with deep classification","1 (ETH)","2 (+SOL)","3 (+BTC)","5+"],
    ["Solana tokens tracked","7","20+","30+","50+"],
    ["Classified transactions in DB","~100K","200K","500K","1M"],
    ["Twitter influencers tracked","0","0","50+","100+"],
    ["Active user alerts","0","50","200","500"],
    ["Grant funding received","$0","$0","$10-15K","$15-25K"],
    ["Blog posts","20","25","40","60"],
])

# ═══ 8. AGENDA ═══
doc.add_heading("Call Agenda", level=1)
table(["Item","Time","Notes"],[
    ["Pep talk","5 min","Weekly sprints. 10-15 hrs/week minimum. Ship every Friday."],
    ["Saif's update","10 min","What shipped, what's blocking, what's next."],
    ["Week 1 sprint planning","15 min","Assign tasks. Dashboard fix + Solana tokens + CryptoPanic."],
    ["Product vision","10 min","Demo Arkham. Discuss Twitter tracking. Alerts priority."],
    ["Eduardo's update","5 min","Grants, marketing, events."],
    ["Feature brainstorm","10 min","What makes a trader open Sonar every single day?"],
    ["Action items","5 min","Who does what, done by Friday."],
])

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
